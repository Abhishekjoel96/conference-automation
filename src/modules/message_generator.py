import os
import json
import logging
import tempfile
from typing import Dict, List, Any, Optional
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Import services
from ..services.openai_service import get_openai_service
from ..services.google_drive import get_drive_service
from ..services.serpapi import get_serpapi_service

# Import other modules
from .linkedin_research import get_linkedin_researcher

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class MessageGenerator:
    def __init__(self):
        self.openai = get_openai_service()
        self.drive_service = get_drive_service()
        self.linkedin_researcher = get_linkedin_researcher()
    
    def generate_message_for_participant(self,
                                      event_name: str,
                                      participant_data: Dict[str, Any],
                                      user_name: str,
                                      user_role: str,
                                      user_company_name: str,
                                      user_company_description: str) -> Optional[Dict[str, Any]]:
        """
        Research a participant and generate a personalized message.
        
        Args:
            event_name: Name of the event
            participant_data: Dictionary with participant information
            user_name: Name of the user
            user_role: Role/position of the user
            user_company_name: Name of the user's company
            user_company_description: Description of the user's company
            linkedin_phantom_id: Phantombuster phantom ID for LinkedIn scraping
            
        Returns:
            Dictionary with message information or None if generation failed
        """
        logger.info(f"Generating message for {participant_data.get('name')} from {participant_data.get('company')}")
        
        # Step 1: Research the participant
        research_data = self.linkedin_researcher.research_participant(
            participant_name=participant_data.get('name', ''),
            participant_role=participant_data.get('role', ''),
            participant_company=participant_data.get('company', ''),
            linkedin_url=participant_data.get('linkedin_url', None),
            # No phantom ID needed for ProxyCurl,
            user_company_name=user_company_name,
            additional_context=participant_data.get('notes', '')
        )
        
        logger.info(f"Research completed for {participant_data.get('name')}. Generating research summary...")
        
        
        # Generate research summary
        research_summary = self.linkedin_researcher.generate_research_summary(research_data)
        
        # Step 2: Identify synergies
        try:
            synergy_points = self.openai.identify_synergies(
                research_summary=research_summary,
                user_company_description=user_company_description
            )
            logger.info(f"Successfully identified synergies for {participant_data.get('name')}")
        except Exception as e:
            logger.error(f"Error identifying synergies: {e}")
            synergy_points = []
        
        # Step 3: Generate personalized message
        try:
            message_data = self.openai.generate_personalized_message(
                participant_name=participant_data.get('name', ''),
                participant_role=participant_data.get('role', ''),
                participant_company=participant_data.get('company', ''),
                research_summary=research_summary,
                synergy_points=synergy_points,
                user_name=user_name,
                user_role=user_role,
                user_company_name=user_company_name,
                user_company_description=user_company_description
            )
            logger.info(f"Successfully generated message for {participant_data.get('name')}")
        except Exception as e:
            logger.error(f"Error generating message: {e}")
            return None
        
        # Step 4: Create and save the message document
        try:
            # Get event folders
            event_folders = self.drive_service.create_event_folders(event_name)
            outreach_folder_id = event_folders.get(f"Outreach_Drafts_{event_name}")
            
            if not outreach_folder_id:
                logger.error(f"Could not find outreach folder for {event_name}")
                return None
            
            # Create docx file
            docx_path = self._create_message_docx(
                participant_name=participant_data.get('name', ''),
                message_data=message_data
            )
            
            # Upload to Google Drive
            file_name = f"{participant_data.get('name', 'Unknown')} - Draft Message.docx"
            
            logger.info(f"Uploading message document to Google Drive for {participant_data.get('name')}")
            file_id = self.drive_service.upload_docx_to_drive(
                file_path=docx_path,
                folder_id=outreach_folder_id,
                file_name=file_name
            )
            
            # Also create a PDF version for better viewability
            pdf_path = self._convert_docx_to_pdf(docx_path)
            if pdf_path:
                pdf_file_name = f"{participant_data.get('name', 'Unknown')} - Draft Message.pdf"
                pdf_file_id = self.drive_service.upload_file_to_drive(
                    file_path=pdf_path,
                    folder_id=outreach_folder_id,
                    file_name=pdf_file_name,
                    mime_type='application/pdf'
                )
                # Delete temporary PDF file
                os.remove(pdf_path)
            else:
                pdf_file_id = None
                logger.warning(f"Could not create PDF for {participant_data.get('name')}")
            
            # Delete temporary DOCX file
            os.remove(docx_path)
            
            logger.info(f"Successfully saved message documents for {participant_data.get('name')}")
            
            # Track in approval tracker
            try:
                from ..modules.approval_tracker import get_approval_tracker
                approval_tracker = get_approval_tracker()
                approval_tracker.add_message(
                    event_name=event_name,
                    participant_name=participant_data.get('name', ''),
                    message_data=message_data,
                    file_id=file_id
                )
                logger.info(f"Added message to approval tracker for {participant_data.get('name')}")
            except Exception as e:
                logger.error(f"Error adding message to approval tracker: {e}")
        except Exception as e:
            logger.error(f"Error creating or uploading document: {e}")
            file_id = None
            pdf_file_id = None
            
            # Return message data with file info
            return {
                "participant": participant_data,
                "research_summary": research_summary,
                "synergy_points": synergy_points,
                "message_data": message_data,
                "file": {
                    "id": file_id,
                    "name": file_name,
                    "folder_id": outreach_folder_id
                }
            }
        except Exception as e:
            logger.error(f"Error creating or saving message document: {e}")
            return None
    
    def _convert_docx_to_pdf(self, docx_path: str) -> Optional[str]:
        """
        Convert a DOCX file to PDF using an external library.
        
        Args:
            docx_path: Path to the DOCX file
            
        Returns:
            Path to the generated PDF file or None if conversion failed
        """
        try:
            # First try using python-docx2pdf if available
            try:
                from docx2pdf import convert
                pdf_path = docx_path.replace('.docx', '.pdf')
                convert(docx_path, pdf_path)
                return pdf_path
            except ImportError:
                logger.warning("docx2pdf not available, trying alternative conversion method")
            
            # Alternative: just return docx path for now as PDF isn't critical
            # In production, we would implement a robust conversion method here
            logger.warning("PDF conversion not implemented, using DOCX only")
            return None
            
        except Exception as e:
            logger.error(f"Error converting DOCX to PDF: {e}")
            return None
    
    def _create_message_docx(self, participant_name: str, message_data: Dict[str, Any]) -> str:
        """
        Create a formatted docx file with the message data following the exact format specified.
        
        Args:
            participant_name: Name of the participant
            message_data: Dictionary with message data
            
        Returns:
            File path to the created docx file
        """
        # Create a new Document
        doc = Document()
        
        # Set default font
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        
        # Add Research Summary section
        doc.add_heading("Research Summary:", level=2)
        
        # Format research summary as key-value pairs
        research_summary = message_data.get("research_summary", {})
        
        # Name
        p = doc.add_paragraph()
        p.add_run("Name: ").bold = True
        p.add_run(research_summary.get("name", participant_name))
        
        # Role
        p = doc.add_paragraph()
        p.add_run("Role: ").bold = True
        p.add_run(research_summary.get("role", ""))
        
        # Company
        p = doc.add_paragraph()
        p.add_run("Company: ").bold = True
        p.add_run(research_summary.get("company", ""))
        
        # LinkedIn
        p = doc.add_paragraph()
        p.add_run("LinkedIn: ").bold = True
        p.add_run(research_summary.get("linkedin", "LinkedIn Profile"))
        
        # Background
        p = doc.add_paragraph()
        p.add_run("Background: ").bold = True
        p.add_run(research_summary.get("background", ""))
        
        # Areas of Synergy
        doc.add_heading("Areas of Synergy:", level=2)
        synergy_areas = research_summary.get("areas_of_synergy", [])
        for point in synergy_areas:
            p = doc.add_paragraph()
            p.add_run(point)
        
        # Add Message Draft section
        doc.add_heading("Message Draft:", level=2)
        message_draft = message_data.get("message_draft", "")
        p = doc.add_paragraph(message_draft)
        
        # Save to temporary file
        temp_dir = tempfile.gettempdir()
        file_path = os.path.join(temp_dir, f"{participant_name} - Draft Message.docx")
        doc.save(file_path)
        
        return file_path
    
    def process_participants_batch(self,
                                event_name: str,
                                participants_data: List[Dict[str, Any]],
                                user_name: str,
                                user_role: str,
                                user_company_name: str,
                                user_company_description: str,
                                ) -> Dict[str, Any]:
        """
        Process a batch of participants to generate personalized messages.
        
        Args:
            event_name: Name of the event
            participants_data: List of participant data dictionaries
            user_name: Name of the user
            user_role: Role/position of the user
            user_company_name: Name of the user's company
            user_company_description: Description of the user's company
            linkedin_phantom_id: Phantombuster phantom ID for LinkedIn scraping
            
        Returns:
            Dictionary with processing statistics
        """
        logger.info(f"Processing batch of {len(participants_data)} participants for {event_name}")
        
        results = {
            "total": len(participants_data),
            "successful": 0,
            "failed": 0,
            "messages": []
        }
        
        # Create approval tracker if it doesn't exist
        approval_tracker_id = self.drive_service.create_approval_tracker(event_name)
        
        # Process each participant
        for participant_data in participants_data:
            try:
                message_result = self.generate_message_for_participant(
                    event_name=event_name,
                    participant_data=participant_data,
                    user_name=user_name,
                    user_role=user_role,
                    user_company_name=user_company_name,
                    user_company_description=user_company_description,
                    # No phantom ID needed for ProxyCurl
                )
                
                if message_result:
                    # Add to results
                    results["successful"] += 1
                    results["messages"].append({
                        "participant": participant_data.get("name"),
                        "status": "success",
                        "file_id": message_result["file"]["id"]
                    })
                    
                    # Update approval tracker
                    self.drive_service.add_participant_to_approval_tracker(
                        spreadsheet_id=approval_tracker_id,
                        participant_name=participant_data.get("name", ""),
                        company_name=participant_data.get("company", "")
                    )
                else:
                    results["failed"] += 1
                    results["messages"].append({
                        "participant": participant_data.get("name"),
                        "status": "failed",
                        "error": "Failed to generate message"
                    })
            except Exception as e:
                logger.error(f"Error processing participant {participant_data.get('name')}: {e}")
                results["failed"] += 1
                results["messages"].append({
                    "participant": participant_data.get("name"),
                    "status": "failed",
                    "error": str(e)
                })
        
        logger.info(f"Batch processing complete. Successful: {results['successful']}, Failed: {results['failed']}")
        return results

# Initialize the message generator
def get_message_generator():
    """Get an instance of the message generator."""
    return MessageGenerator()
